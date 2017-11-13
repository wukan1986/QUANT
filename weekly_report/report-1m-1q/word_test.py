# -*- coding: utf-8 -*-
"""
Created on Tue Sep 05 14:46:52 2017

@author: shiyunchao
"""

import pandas as pd
from time import strftime, localtime, time
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor

document = Document()
style = document.styles['Normal']
font = style.font

document.styles['Normal'].font.name = u'微软雅黑'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

from docx.shared import Pt
#font.name = u'宋体'
font.size = Pt(8)

#a = document.add_heading(u'Barra模型因子分析', 0)
#a.style.font.name = u'微软雅黑'
#a.style._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

document.add_paragraph(u'风险模型因子分析','Title')

p1 = document.add_paragraph(u'基于Barra框架，分析各种风险因子，对其表现按日频率跟踪','Body Text 2')
p1.style.font.size = Pt(9.5)

#p.add_run(u'测试').bold = True
#p.add_run(u' 和一些 ')
#p.add_run(u'字体').italic = True

#document.add_heading(u'风格因子评价指标跟踪', level=1)
#p2 = document.add_paragraph('                     ', style='IntenseQuote')
#p2.style.font.size = Pt(1)

p3 = document.add_paragraph(u'风格因子评价指标跟踪', style='ListBullet')
#p3.style.font.size = Pt(14)
#p3.style.font.color.rgb = RGBColor(85, 115, 145)

#document.add_paragraph(
#    'first item in unordered list', style='ListBullet'
#)
p4 = document.add_paragraph(u'风格因子累计收益率 最近一个月', style='ListNumber')
p4.style.font.size = Pt(10)

document.add_picture('data/factor_month.png', width=Inches(6))

document.add_paragraph(u'风格因子累计收益率 最近一个季度', style='ListNumber')
document.add_picture('data/factor_year.png', width=Inches(6))

document.add_paragraph(' ')

document.add_paragraph(u'申万行业因子累计收益率 最近一个月', style='ListNumber')
document.add_picture('data/industry_month.png', width=Inches(6))

document.add_paragraph(' ')

document.add_paragraph(u'申万行业因子累计收益率 最近一个季度', style='ListNumber')
document.add_picture('data/industry_year.png', width=Inches(6))

document.add_paragraph(' ')
document.add_paragraph(' ')
document.add_paragraph(' ')

p3 = document.add_paragraph(u'评价', style='ListBullet')
p3.style.font.size = Pt(14)
p3.style.font.color.rgb = RGBColor(85, 115, 145)

f = open('data/comment.txt')
comment = f.read().decode('utf-8')
f.close()

document.add_paragraph(u'%s'%comment,'Body Text 2')

document.add_paragraph(' ')

p3 = document.add_paragraph(u'风格因子收益与风险分析', style='ListBullet')
p3.style.font.size = Pt(14)
p3.style.font.color.rgb = RGBColor(85, 115, 145)

document.add_paragraph(u'1、 风格因子 最近一个月','Body Text 2')

table = document.add_table(rows=1, cols=9)
table.style.font.size = Pt(8)
hdr_cells = table.rows[0].cells
col = ['FACTOR','Return', 'MDD','MDD_SPAN','Sharpe','Daily_Min','Min_Day','Win_Ratio','P2L_Ratio']
col = [u'因子',u'收益率', u'最大回撤',u'最大回测区间',u'夏普比率',u'日收益最小值',u'日收益最小日期',u'胜率',u'盈亏比率']
for i in range(len(col)):
    hdr_cells[i].text = col[i]



recordset = pd.read_csv('data/indicator_factor.csv', encoding = 'gb18030',index_col = 0)


for i,item in recordset.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = item.name
    row_cells[1].text = str(item.Return)
    row_cells[2].text = str(item.MDD)
    row_cells[3].text = str(item.MDD_SPAN)
    row_cells[4].text = str(item.Sharpe)
    row_cells[5].text = str(item.Daily_Min)
    row_cells[6].text = str(item.Min_Day)
    row_cells[7].text = str(item.Win_Ratio)
    row_cells[8].text = str(item.P2L_Ratio)

table.style = 'LightShading-Accent1'

document.add_paragraph(' ')
document.add_paragraph(u'2、 申万行业因子 最近一个月','Body Text 2')

table2 = document.add_table(rows=1, cols=9)
table2.style.font.size = Pt(8)
hdr_cells = table2.rows[0].cells
col = ['FACTOR','Return', 'MDD','MDD_SPAN','Sharpe','Daily_Min','Min_Day','Win_Ratio','P2L_Ratio']
col = [u'因子',u'收益率', u'最大回撤',u'最大回测区间',u'夏普比率',u'日收益最小值',u'日收益最小日期',u'胜率',u'盈亏比率']
for i in range(len(col)):
    hdr_cells[i].text = col[i]


recordset = pd.read_csv('data/indicator_industry.csv', encoding = 'gb18030',index_col = 0)


for i,item in recordset.iterrows():
    row_cells = table2.add_row().cells
    row_cells[0].text = item.name
    row_cells[1].text = str(item.Return)
    row_cells[2].text = str(item.MDD)
    row_cells[3].text = str(item.MDD_SPAN)
    row_cells[4].text = str(item.Sharpe)
    row_cells[5].text = str(item.Daily_Min)
    row_cells[6].text = str(item.Min_Day)
    row_cells[7].text = str(item.Win_Ratio)
    row_cells[8].text = str(item.P2L_Ratio)

table2.style = 'LightShading-Accent1'

#####要添加
document.add_paragraph(' ')
document.add_paragraph(' ')
p5 = document.add_paragraph(u'附录', style='ListBullet')
p5.style.font.size = Pt(14)
p5.style.font.color.rgb = RGBColor(85, 115, 145)

f = open('data/notes.txt')
notes = f.read().decode('utf-8')
f.close()

document.add_paragraph(u'%s'%notes,'Body Text 2')

#document.add_page_break()

today = strftime("%Y%m%d",localtime())
document.save('month_%s.docx'%today)